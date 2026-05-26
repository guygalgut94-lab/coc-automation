import streamlit as st
import pdfplumber
import re
from docx import Document
from docx.shared import Inches
from io import BytesIO
from datetime import datetime
import os

st.title("COC Automation")

uploaded_files = st.file_uploader(
    "Upload COT PDF files",
    type="pdf",
    accept_multiple_files=True
)

worker = st.selectbox(
    "Certified by",
    ["Guy", "Ilan", "Mark", "Conley"]
)

def extract_text_from_pdf(uploaded_file):
    text = ""

    with pdfplumber.open(uploaded_file) as pdf:
        for page in pdf.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"

    return text

def get_value_between(text, start, end):
    pattern = rf"{re.escape(start)}\s*(.*?)\s*{re.escape(end)}"
    match = re.search(pattern, text, re.DOTALL)
    return match.group(1).strip() if match else ""

def extract_item_data(text, file_name):
    supplier = get_value_between(text, "Supplier Name", "Sampling Plan")
    customer = get_value_between(text, "Customer", "Level 2")
    po = get_value_between(text, "Customer ID", "Part Name")

    part_no_match = re.search(r'Part No\.\s*(\S+)', text)
    desc_match = re.search(r'Part Name\s*(.*?)\s+Material', text)
    rev_match = re.search(r'Drawing Rev\s*(\d+)', text)
    qty_match = re.search(r'Lot Qty\s*(\d+)', text)

    return {
        "file_name": file_name,
        "supplier": supplier,
        "customer": customer,
        "po": po,
        "part_no": part_no_match.group(1).strip() if part_no_match else "",
        "description": desc_match.group(1).strip() if desc_match else "",
        "rev": rev_match.group(1).strip() if rev_match else "",
        "qty": qty_match.group(1).strip() if qty_match else ""
    }

def replace_placeholders_in_paragraph(paragraph, replacements):
    for key, value in replacements.items():
        if key in paragraph.text:
            paragraph.text = paragraph.text.replace(key, value)

def replace_placeholders(doc, replacements):
    for paragraph in doc.paragraphs:
        replace_placeholders_in_paragraph(paragraph, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_placeholders_in_paragraph(paragraph, replacements)

def insert_signature(doc, worker_name):
    signature_path = f"signatures/{worker_name}.png"

    if not os.path.exists(signature_path):
        st.warning(f"Signature image not found: {signature_path}")
        return

    for paragraph in doc.paragraphs:
        if "[signature]" in paragraph.text:
            paragraph.text = paragraph.text.replace("[signature]", "")
            run = paragraph.add_run()
            run.add_picture(signature_path, width=Inches(0.9))
            return

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if "[signature]" in paragraph.text:
                        paragraph.text = paragraph.text.replace("[signature]", "")
                        run = paragraph.add_run()
                        run.add_picture(signature_path, width=Inches(0.9))
                        return

    st.warning("[signature] placeholder not found in template.")

if uploaded_files:
    items = []

    for uploaded_file in uploaded_files:
        text = extract_text_from_pdf(uploaded_file)
        item = extract_item_data(text, uploaded_file.name)
        items.append(item)

    coc_date = datetime.today().strftime("%d %b %Y")

    default_supplier = items[0]["supplier"] if items else ""
    default_customer = items[0]["customer"] if items else ""
    default_po = items[0]["po"] if items else ""

    st.subheader("Header Information")

    supplier_name = st.text_input(
        "Supplier Name",
        value=default_supplier
    )

    customer_name = st.text_input(
        "Customer Name",
        value=default_customer
    )

    po_number = st.text_input(
        "PO Number",
        value=default_po
    )

    st.subheader("Extracted Parts")

    st.dataframe([
        {
            "#": index + 1,
            "File Name": item["file_name"],
            "Supplier": item["supplier"],
            "Customer": item["customer"],
            "PO Number": item["po"],
            "Part No": item["part_no"],
            "Description": item["description"],
            "REV": item["rev"],
            "QTY": item["qty"]
        }
        for index, item in enumerate(items)
    ])

    if st.button("Generate COC"):
        doc = Document("COC Format.docx")

        replace_placeholders(doc, {
            "[Date]": coc_date,
            "[Supplier Name]": supplier_name,
            "[Customer Name]": customer_name,
            "[Customer Id]": po_number
        })

        parts_table = doc.tables[1]

        while len(parts_table.rows) > 1:
            row = parts_table.rows[1]
            row._element.getparent().remove(row._element)

        for index, item in enumerate(items, start=1):
            row_cells = parts_table.add_row().cells
            row_cells[0].text = str(index)
            row_cells[1].text = item["part_no"]
            row_cells[2].text = item["description"]
            row_cells[3].text = item["rev"]
            row_cells[4].text = item["qty"]

        insert_signature(doc, worker)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        file_name = f"COC_{po_number or customer_name or 'Generated'}.docx"

        st.download_button(
            label="Download COC DOCX",
            data=buffer,
            file_name=file_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
